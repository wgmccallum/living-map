"""Document parsing for the Moderated Bulk Add ingest pipeline.

Extracts structured content chunks from source documents in supported formats:
.tex, .xlsx, .docx, .txt, .pdf

Each parser returns a list of ContentChunk objects — the raw material for
AI-driven KC decomposition in Stage 1.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ContentChunk:
    """A single content unit extracted from a source document."""
    source_reference: str          # Identifier in the source (e.g., standard ID)
    original_text: str             # Verbatim text from the source
    metadata: dict = field(default_factory=dict)  # Format-specific metadata


# ═══════════════════════════════════════════════════════
# .tex parser (CCSS standards)
# ═══════════════════════════════════════════════════════

_TEX_FIELD_RE = re.compile(r"\\(\w+)\{([^}]*)\}")


def parse_tex(content: str, filename: str = "") -> list[ContentChunk]:
    """Parse a .tex file containing CCSS-style standard definitions.

    Handles both single-standard files and multi-standard files
    (e.g., allstandards.tex or concatenated files).
    """
    chunks = []

    # Split on \standardid to handle multi-standard files
    # First, check if there are multiple standards
    parts = re.split(r"(?=\\standardid\{)", content)

    for part in parts:
        part = part.strip()
        if not part or "\\standardid" not in part:
            continue

        fields = {}
        for match in _TEX_FIELD_RE.finditer(part):
            cmd, value = match.group(1), match.group(2)
            fields[cmd] = value.strip()

        standard_id = fields.get("standardid", "")
        text = fields.get("text", "")

        if not standard_id or not text:
            continue

        metadata = {
            k: v for k, v in fields.items()
            if k not in ("standardid", "text") and v
        }

        chunks.append(ContentChunk(
            source_reference=standard_id,
            original_text=text,
            metadata=metadata,
        ))

    return chunks


def parse_tex_directory(dir_path: Path) -> list[ContentChunk]:
    """Parse all .tex files in a directory."""
    chunks = []
    for tex_file in sorted(dir_path.glob("*.tex")):
        content = tex_file.read_text(encoding="utf-8", errors="replace")
        chunks.extend(parse_tex(content, tex_file.name))
    return chunks


# ═══════════════════════════════════════════════════════
# .xlsx parser
# ═══════════════════════════════════════════════════════

def parse_xlsx(file_path: Path) -> list[ContentChunk]:
    """Parse an Excel spreadsheet into content chunks.

    Uses a heuristic approach: reads all sheets, treats each row as a
    potential content chunk. The first row is assumed to be headers.
    Columns are identified by header names when possible.
    """
    import openpyxl

    wb = openpyxl.load_workbook(str(file_path), data_only=True)
    chunks = []

    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 2:
            continue

        # First row = headers
        headers = [str(h).strip().lower() if h else f"col_{i}" for i, h in enumerate(rows[0])]

        # Try to identify key columns
        id_col = _find_column(headers, ["id", "kc_id", "standard_id", "standardid", "code"])
        text_col = _find_column(headers, ["text", "description", "short_description", "standard_text", "content"])
        detail_col = _find_column(headers, ["long_description", "detail", "notes", "elaboration"])

        for row_idx, row in enumerate(rows[1:], start=2):
            if not any(cell for cell in row):
                continue

            # Build reference
            ref = ""
            if id_col is not None and row[id_col]:
                ref = str(row[id_col]).strip()
            else:
                ref = f"{ws.title}:row{row_idx}"

            # Build text
            text_parts = []
            if text_col is not None and row[text_col]:
                text_parts.append(str(row[text_col]).strip())
            if detail_col is not None and row[detail_col]:
                text_parts.append(str(row[detail_col]).strip())

            # If no recognized columns, concatenate all non-empty cells
            if not text_parts:
                text_parts = [str(cell).strip() for cell in row if cell]

            if not text_parts:
                continue

            original_text = " | ".join(text_parts)

            # Collect all columns as metadata
            metadata = {"sheet": ws.title}
            for i, header in enumerate(headers):
                if i < len(row) and row[i] is not None:
                    val = row[i]
                    if i not in (id_col, text_col, detail_col):
                        metadata[header] = str(val).strip() if not isinstance(val, (int, float)) else val

            chunks.append(ContentChunk(
                source_reference=ref,
                original_text=original_text,
                metadata=metadata,
            ))

    wb.close()
    return chunks


def _find_column(headers: list[str], candidates: list[str]) -> int | None:
    """Find the first header matching any candidate name."""
    for candidate in candidates:
        for i, h in enumerate(headers):
            if candidate in h:
                return i
    return None


# ═══════════════════════════════════════════════════════
# .docx parser
# ═══════════════════════════════════════════════════════

def parse_docx(file_path: Path) -> list[ContentChunk]:
    """Parse a Word document into content chunks.

    Groups content by headings. Each heading starts a new chunk;
    the body paragraphs under it become the chunk's text.
    If no headings are found, each non-empty paragraph is a chunk.
    """
    import docx

    doc = docx.Document(str(file_path))
    chunks = []

    # Check if document uses headings
    has_headings = any(
        p.style.name.startswith("Heading") for p in doc.paragraphs
        if p.style and p.style.name
    )

    if has_headings:
        current_heading = None
        current_body = []
        chunk_num = 0

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            if style_name.startswith("Heading"):
                # Save previous chunk
                if current_heading and current_body:
                    chunk_num += 1
                    chunks.append(ContentChunk(
                        source_reference=f"section-{chunk_num}",
                        original_text="\n".join(current_body),
                        metadata={"heading": current_heading, "heading_style": style_name},
                    ))
                current_heading = text
                current_body = []
            else:
                current_body.append(text)

        # Don't forget last chunk
        if current_heading and current_body:
            chunk_num += 1
            chunks.append(ContentChunk(
                source_reference=f"section-{chunk_num}",
                original_text="\n".join(current_body),
                metadata={"heading": current_heading},
            ))
    else:
        # No headings — each paragraph is a chunk
        for i, para in enumerate(doc.paragraphs, start=1):
            text = para.text.strip()
            if text:
                chunks.append(ContentChunk(
                    source_reference=f"para-{i}",
                    original_text=text,
                    metadata={},
                ))

    # Also extract tables
    for t_idx, table in enumerate(doc.tables, start=1):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            chunks.append(ContentChunk(
                source_reference=f"table-{t_idx}",
                original_text="\n".join(rows_text),
                metadata={"type": "table"},
            ))

    return chunks


# ═══════════════════════════════════════════════════════
# .pdf parser
# ═══════════════════════════════════════════════════════

def parse_pdf(file_path: Path) -> list[ContentChunk]:
    """Parse a PDF into content chunks.

    Extracts text page-by-page.  Consecutive pages whose text forms a
    single logical block (no blank-line break) are merged.  Tables
    detected by pdfplumber are appended as separate chunks.
    """
    import pdfplumber

    chunks = []

    with pdfplumber.open(str(file_path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                chunks.append(ContentChunk(
                    source_reference=f"page-{page_num}",
                    original_text=text,
                    metadata={"page": page_num},
                ))

            # Extract tables separately
            for t_idx, table in enumerate(page.extract_tables() or [], start=1):
                rows_text = []
                for row in table:
                    cells = [str(cell).strip() if cell else "" for cell in row]
                    rows_text.append(" | ".join(cells))
                if rows_text:
                    chunks.append(ContentChunk(
                        source_reference=f"page-{page_num}-table-{t_idx}",
                        original_text="\n".join(rows_text),
                        metadata={"page": page_num, "type": "table"},
                    ))

    return chunks


# ═══════════════════════════════════════════════════════
# Plain text parser
# ═══════════════════════════════════════════════════════

def parse_text(content: str, filename: str = "") -> list[ContentChunk]:
    """Parse plain text into content chunks.

    Splits on blank lines (double newlines) to identify paragraphs.
    Each non-empty paragraph becomes a chunk.
    """
    chunks = []
    paragraphs = re.split(r"\n\s*\n", content)

    for i, para in enumerate(paragraphs, start=1):
        text = para.strip()
        if text:
            chunks.append(ContentChunk(
                source_reference=f"chunk-{i}",
                original_text=text,
                metadata={},
            ))

    return chunks


# ═══════════════════════════════════════════════════════
# Unified parser dispatch
# ═══════════════════════════════════════════════════════

def parse_document(file_path: Path, content: bytes | None = None) -> list[ContentChunk]:
    """Parse a document file into content chunks.

    Args:
        file_path: Path to the file (used for format detection and reading)
        content: Optional raw bytes (if the file is uploaded and not yet on disk)

    Returns:
        List of ContentChunk objects
    """
    suffix = file_path.suffix.lower()

    if suffix == ".tex":
        if content:
            text = content.decode("utf-8", errors="replace")
        else:
            text = file_path.read_text(encoding="utf-8", errors="replace")
        return parse_tex(text, file_path.name)

    elif suffix == ".xlsx":
        # openpyxl needs a file on disk
        if content:
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
                tmp.write(content)
                tmp.flush()
                return parse_xlsx(Path(tmp.name))
        return parse_xlsx(file_path)

    elif suffix == ".docx":
        if content:
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(content)
                tmp.flush()
                return parse_docx(Path(tmp.name))
        return parse_docx(file_path)

    elif suffix == ".pdf":
        if content:
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(content)
                tmp.flush()
                return parse_pdf(Path(tmp.name))
        return parse_pdf(file_path)

    elif suffix in (".txt", ".text", ".md"):
        if content:
            text = content.decode("utf-8", errors="replace")
        else:
            text = file_path.read_text(encoding="utf-8", errors="replace")
        return parse_text(text, file_path.name)

    else:
        raise ValueError(f"Unsupported file format: {suffix}")
