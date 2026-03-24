"""Tests for the document ingest pipeline (parser + endpoint)."""

import io
import tempfile
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from living_map.app import app
from living_map.document_parser import (
    ContentChunk,
    parse_document,
    parse_tex,
    parse_text,
)


@pytest.fixture
def client(tmp_path):
    db_path = tmp_path / "test.db"
    app.state.db_path = str(db_path)
    with TestClient(app) as c:
        yield c


# ═══════════════════════════════════════════════════════
# Document Parser Unit Tests
# ═══════════════════════════════════════════════════════


class TestTexParser:
    def test_single_standard(self):
        content = r"""
\standardid{K.CC.1}
\document{Math}
\grade{K}
\domain{Counting and Cardinality}
\domainid{CC}
\cluster{Know number names and the count sequence.}
\standardnumber{1}
\text{Count to 100 by ones and by tens.}
"""
        chunks = parse_tex(content)
        assert len(chunks) == 1
        assert chunks[0].source_reference == "K.CC.1"
        assert chunks[0].original_text == "Count to 100 by ones and by tens."
        assert chunks[0].metadata["grade"] == "K"
        assert chunks[0].metadata["domain"] == "Counting and Cardinality"

    def test_multiple_standards(self):
        content = r"""
\standardid{K.CC.1}
\grade{K}
\text{Count to 100 by ones and by tens.}

\standardid{K.CC.2}
\grade{K}
\text{Count forward beginning from a given number.}
"""
        chunks = parse_tex(content)
        assert len(chunks) == 2
        assert chunks[0].source_reference == "K.CC.1"
        assert chunks[1].source_reference == "K.CC.2"

    def test_empty_text_skipped(self):
        content = r"""
\standardid{K.CC.1}
\text{}
"""
        chunks = parse_tex(content)
        assert len(chunks) == 0


class TestTextParser:
    def test_paragraph_splitting(self):
        content = """First paragraph about counting.

Second paragraph about addition.

Third paragraph about subtraction."""
        chunks = parse_text(content)
        assert len(chunks) == 3
        assert "counting" in chunks[0].original_text
        assert "addition" in chunks[1].original_text

    def test_empty_content(self):
        chunks = parse_text("")
        assert len(chunks) == 0


class TestXlsxParser:
    def test_basic_spreadsheet(self, tmp_path):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Standards"
        ws.append(["id", "text", "grade"])
        ws.append(["K.CC.1", "Count to 100", "K"])
        ws.append(["K.CC.2", "Count forward", "K"])
        path = tmp_path / "test.xlsx"
        wb.save(str(path))

        chunks = parse_document(path)
        assert len(chunks) == 2
        assert chunks[0].source_reference == "K.CC.1"
        assert "Count to 100" in chunks[0].original_text


class TestDocxParser:
    def test_heading_based_parsing(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_heading("Section One", level=1)
        doc.add_paragraph("Content under section one.")
        doc.add_paragraph("More content here.")
        doc.add_heading("Section Two", level=1)
        doc.add_paragraph("Content under section two.")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        chunks = parse_document(path)
        assert len(chunks) == 2
        assert "Content under section one" in chunks[0].original_text
        assert chunks[0].metadata["heading"] == "Section One"

    def test_no_headings(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        chunks = parse_document(path)
        assert len(chunks) == 2


class TestUnifiedParser:
    def test_unsupported_format(self, tmp_path):
        path = tmp_path / "test.xyz"
        path.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported"):
            parse_document(path)

    def test_tex_via_content_bytes(self):
        content = rb"""
\standardid{K.CC.1}
\text{Count to 100 by ones and by tens.}
"""
        chunks = parse_document(Path("test.tex"), content=content)
        assert len(chunks) == 1

    def test_text_via_content_bytes(self):
        content = b"Para one.\n\nPara two."
        chunks = parse_document(Path("test.txt"), content=content)
        assert len(chunks) == 2


# ═══════════════════════════════════════════════════════
# Ingest Endpoint Tests
# ═══════════════════════════════════════════════════════


class TestIngestEndpoint:
    def _create_session(self, client):
        client.post("/api/staging", json={
            "id": "test-ingest",
            "topic_name": "Ingest Test",
        })

    def test_ingest_tex_file(self, client):
        self._create_session(client)
        content = rb"""
\standardid{K.CC.1}
\grade{K}
\domain{Counting and Cardinality}
\text{Count to 100 by ones and by tens.}

\standardid{K.CC.2}
\grade{K}
\domain{Counting and Cardinality}
\text{Count forward beginning from a given number.}

\standardid{K.CC.3}
\grade{K}
\domain{Counting and Cardinality}
\text{Write numbers from 0 to 20.}
"""
        resp = client.post(
            "/api/staging/test-ingest/ingest",
            files={"file": ("standards.tex", io.BytesIO(content), "application/x-tex")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["session_id"] == "test-ingest"
        assert data["chunks_extracted"] == 3
        assert data["staged_kcs_created"] == 3
        assert data["format"] == ".tex"

        # Verify manifest
        manifest = data["manifest"]
        assert len(manifest) == 3
        assert manifest[0]["source_reference"] == "K.CC.1"
        assert manifest[0]["original_text"] == "Count to 100 by ones and by tens."

        # Verify staged KCs were created
        resp = client.get("/api/staging/test-ingest/kcs")
        kcs = resp.json()
        assert len(kcs) == 3
        assert all(kc["stage_status"] == "proposed" for kc in kcs)
        assert kcs[0]["source_reference"] == "K.CC.1"

    def test_ingest_text_file(self, client):
        self._create_session(client)
        content = b"Students should count to 10.\n\nStudents add single digits."
        resp = client.post(
            "/api/staging/test-ingest/ingest",
            files={"file": ("content.txt", io.BytesIO(content), "text/plain")},
        )
        assert resp.status_code == 200
        assert resp.json()["chunks_extracted"] == 2

    def test_ingest_updates_source_documents(self, client):
        self._create_session(client)
        content = rb"\standardid{K.CC.1}\text{Count to 100.}"
        client.post(
            "/api/staging/test-ingest/ingest",
            files={"file": ("my_standards.tex", io.BytesIO(content), "application/x-tex")},
        )
        resp = client.get("/api/staging/test-ingest")
        assert "my_standards.tex" in resp.json()["source_documents"]

    def test_ingest_nonexistent_session(self, client):
        content = rb"\standardid{K.CC.1}\text{Count to 100.}"
        resp = client.post(
            "/api/staging/nonexistent/ingest",
            files={"file": ("test.tex", io.BytesIO(content), "application/x-tex")},
        )
        assert resp.status_code == 404

    def test_ingest_unsupported_format(self, client):
        self._create_session(client)
        resp = client.post(
            "/api/staging/test-ingest/ingest",
            files={"file": ("test.xyz", io.BytesIO(b"hello"), "application/octet-stream")},
        )
        assert resp.status_code == 400

    def test_ingest_empty_file(self, client):
        self._create_session(client)
        resp = client.post(
            "/api/staging/test-ingest/ingest",
            files={"file": ("empty.txt", io.BytesIO(b""), "text/plain")},
        )
        assert resp.status_code == 422

    def test_incremental_ingest(self, client):
        """Ingesting a second document continues numbering from where the first left off."""
        self._create_session(client)

        # First ingest: 2 chunks
        content1 = b"Para one.\n\nPara two."
        client.post(
            "/api/staging/test-ingest/ingest",
            files={"file": ("doc1.txt", io.BytesIO(content1), "text/plain")},
        )

        # Second ingest: 2 more chunks
        content2 = b"Para three.\n\nPara four."
        resp = client.post(
            "/api/staging/test-ingest/ingest",
            files={"file": ("doc2.txt", io.BytesIO(content2), "text/plain")},
        )
        data = resp.json()
        assert data["staged_kcs_created"] == 2

        # Verify all 4 KCs exist with unique IDs
        resp = client.get("/api/staging/test-ingest/kcs")
        kcs = resp.json()
        assert len(kcs) == 4
        ids = [kc["id"] for kc in kcs]
        assert len(set(ids)) == 4  # All unique

    def test_ingest_xlsx(self, client, tmp_path):
        """Test ingesting an Excel file via the endpoint."""
        import openpyxl
        self._create_session(client)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Standards"
        ws.append(["id", "text", "grade"])
        ws.append(["K.CC.1", "Count to 100", "K"])
        ws.append(["K.CC.2", "Count forward", "K"])
        path = tmp_path / "test.xlsx"
        wb.save(str(path))

        with open(path, "rb") as f:
            resp = client.post(
                "/api/staging/test-ingest/ingest",
                files={"file": ("test.xlsx", f, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )
        assert resp.status_code == 200
        assert resp.json()["chunks_extracted"] == 2

    def test_ingest_real_tex_standards(self, client):
        """Test with content matching the actual CCSS .tex format."""
        self._create_session(client)
        content = rb"""
\standardid{3.NF.1}
\document{Math}
\sectionid{Math K-8}
\anchor{0}
\theme{}
\domain{Number and Operations---Fractions}
\domainid{NF}
\grade{3}
\cluster{Develop understanding of fractions as numbers.}
\standardnumber{1}
\letter{}
\adv{}
\mod{}
\text{Understand a fraction 1/b as the quantity formed by 1 part when a whole is partitioned into b equal parts; understand a fraction a/b as the quantity formed by a parts of size 1/b.}

%%% Local Variables:
%%% mode: latex
%%% TeX-master: "allstandards"
%%% End:
"""
        resp = client.post(
            "/api/staging/test-ingest/ingest",
            files={"file": ("3.NF.1.tex", io.BytesIO(content), "application/x-tex")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["chunks_extracted"] == 1
        manifest = data["manifest"]
        assert manifest[0]["source_reference"] == "3.NF.1"
        assert "fraction" in manifest[0]["original_text"].lower()
        assert manifest[0]["metadata"]["grade"] == "3"
        assert manifest[0]["metadata"]["domain"] == "Number and Operations---Fractions"
